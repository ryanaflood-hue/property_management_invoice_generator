import os
import io
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt
from models import Invoice, SessionLocal, Customer, Settings

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "invoice_templates")
TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "base_invoice_template.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "generated_invoices")

os.makedirs(OUTPUT_DIR, exist_ok=True)

def get_invoice_templates():
    """Return a list of available invoice template filenames (docx) in the invoice_templates folder."""
    templates = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx") and not f.startswith("~")]
    return sorted(templates)

def get_period_dates(invoice_date: date, cadence: str):
    """Calculate start and end dates for the period based on cadence."""
    if cadence == "monthly":
        start_date = invoice_date.replace(day=1)
        # End of month calculation
        if invoice_date.month == 12:
            end_date = invoice_date.replace(day=31)
        else:
            next_month = invoice_date.replace(day=28) + timedelta(days=4)
            end_date = next_month - timedelta(days=next_month.day)
    elif cadence == "quarterly":
        quarter = (invoice_date.month - 1) // 3 + 1
        start_month = 3 * (quarter - 1) + 1
        start_date = invoice_date.replace(month=start_month, day=1)
        if start_month + 2 > 12:
            end_date = invoice_date.replace(month=12, day=31)
        else:
            end_month = start_month + 2
            next_month = invoice_date.replace(month=end_month, day=28) + timedelta(days=4)
            end_date = next_month - timedelta(days=next_month.day)
    elif cadence == "yearly":
        start_date = invoice_date.replace(month=1, day=1)
        end_date = invoice_date.replace(month=12, day=31)
    else:
        start_date = invoice_date
        end_date = invoice_date
    
    return start_date, end_date

def get_period_label(invoice_date: date, cadence: str) -> str:
    year = invoice_date.year
    if cadence == "monthly":
        return invoice_date.strftime("%B %Y")  # "March 2025"
    elif cadence == "quarterly":
        quarter = (invoice_date.month - 1) // 3 + 1
        return f"{quarter}rd quarter {year}" if quarter == 3 else f"{quarter}st quarter {year}" if quarter == 1 else f"{quarter}nd quarter {year}" if quarter == 2 else f"{quarter}th quarter {year}"
    elif cadence == "yearly":
        return f"{year}"
    else:
        return invoice_date.isoformat()

def fill_invoice_template(doc, replacements):
    """Replace placeholders in the document with values from replacements dict."""
    # Define keys that need tight spacing
    tight_spacing_keys = ["{{FEE_LINE_2}}", "{{FEE_LINE_3}}", "{{ADDITIONAL_FEE_LINE}}"]

    for p in doc.paragraphs:
        replaced = False
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, str(new))
                replaced = True
                # Apply standard spacing (12pt) if it's a fee line
                if old in tight_spacing_keys:
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.line_spacing = 1.0
        if replaced:
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replaced = False
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, str(new))
                            replaced = True
                            # Apply standard spacing (12pt) if it's a fee line
                            if old in tight_spacing_keys:
                                p.paragraph_format.space_after = Pt(12)
                                p.paragraph_format.line_spacing = 1.0
                    if replaced:
                        for run in p.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(14)

def _generate_invoice_logic(customer, invoice_date, period_label, period_dates, amount, return_buffer=True, **kwargs):
    """
    Shared logic to generate an invoice.
    If return_buffer is True, returns (filename, BytesIO_object).
    If return_buffer is False, saves to file and returns (filename, full_path).
    DEFAULT IS TRUE FOR VERCEL CLOUD COMPATIBILITY (read-only filesystem).
    
    kwargs can contain:
    - fee_2_type, fee_2_amount
    - fee_3_type, fee_3_amount
    - additional_fee_desc, additional_fee_amount
    """
    try:
        doc = Document(TEMPLATE_PATH)
        
        print(f"DEBUG: Generator Logic Called. Kwargs: {kwargs}")
        if kwargs:
            print(f"DEBUG: Using kwargs. fee_2_amount={kwargs.get('fee_2_amount')}, additional={kwargs.get('additional_fee_amount')}")
        
        if kwargs:
            # Manual generation: use provided values (even if None)
            fee_2_type = kwargs.get('fee_2_type')
            fee_2_amount = kwargs.get('fee_2_amount')
            fee_3_type = kwargs.get('fee_3_type')
            fee_3_amount = kwargs.get('fee_3_amount')
            additional_fee_desc = kwargs.get('additional_fee_desc')
            additional_fee_amount = kwargs.get('additional_fee_amount')
        else:
            # Batch generation: use customer defaults
            fee_2_type = customer.fee_2_type
            fee_2_amount = customer.fee_2_rate
            fee_3_type = customer.fee_3_type
            fee_3_amount = customer.fee_3_rate
            additional_fee_desc = customer.additional_fee_desc
            additional_fee_amount = customer.additional_fee_amount
        
        # Calculate total amount including all fees
        # Start with base rate
        total_amount = amount
        
        # Add Fee 2
        if fee_2_amount:
            total_amount += fee_2_amount
            
        # Add Fee 3
        if fee_3_amount:
            total_amount += fee_3_amount
            
        # Add Additional Fee
        if additional_fee_amount:
            total_amount += additional_fee_amount
            
        # Add Property Fees
        property_fees_total = 0
        for prop in customer.properties:
            if prop.fee_amount:
                property_fees_total += prop.fee_amount
        total_amount += property_fees_total
        
        
        # Build complete fee lines (or empty strings if not used)
        # Calculate period info for fee 2 and 3 if they exist
        # Build complete fee lines (or empty strings if not used)
        # Calculate period info for fee 2 and 3 if they exist
        fee_line_2 = ""
        if fee_2_amount:
            # Fallback to "Fee" if type is missing
            f2_type = fee_2_type or "Fee"
            fee_line_2 = f"{period_label} {f2_type} ({period_dates}) = ${fee_2_amount:,.2f}"
        
        fee_line_3 = ""
        if fee_3_amount:
            f3_type = fee_3_type or "Fee"
            fee_line_3 = f"{period_label} {f3_type} ({period_dates}) = ${fee_3_amount:,.2f}"
        
        # Build additional fee line
        additional_fee_parts = []
        if additional_fee_amount:
             additional_fee_parts.append(f"{additional_fee_desc} = ${additional_fee_amount:,.2f}")
        
        # Append property fees
        if customer.properties:
            for prop in customer.properties:
                if prop.fee_amount:
                    additional_fee_parts.append(f"Management Fee ({prop.address}) = ${prop.fee_amount:,.2f}")
        
        additional_fee_line = "\n\n".join(additional_fee_parts)
        
        replacements = {
            "{{CUSTOMER_NAME}}": customer.name,
            "{{CUSTOMER_EMAIL}}": customer.email,
            "{{PROPERTY_ADDRESS}}": customer.property_address,
            "{{PROPERTY_CITY}}": customer.property_city or "",
            "{{PROPERTY_STATE}}": customer.property_state or "",
            "{{PROPERTY_ZIP}}": customer.property_zip or "",
            "{{PERIOD}}": period_label,
            "{{PERIOD_DATES}}": period_dates,
            "{{AMOUNT}}": f"${amount:,.2f}",
            "{{INVOICE_DATE}}": invoice_date.strftime("%m/%d/%Y"),
            "{{FEE_TYPE}}": getattr(customer, "fee_type", "Management Fee") or "Management Fee",
            "{{TOTAL_AMOUNT}}": f"${total_amount:,.2f}",
            # Complete fee lines - these replace the entire row content
            "{{FEE_LINE_2}}": fee_line_2,
            "{{FEE_LINE_3}}": fee_line_3,
            "{{ADDITIONAL_FEE_LINE}}": additional_fee_line,
        }

        # Remove rows with empty fee lines to eliminate whitespace
        rows_to_remove = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Check if this row contains an empty fee line placeholder
                    if ("{{FEE_LINE_2}}" in cell.text and not fee_line_2) or \
                       ("{{FEE_LINE_3}}" in cell.text and not fee_line_3) or \
                       ("{{ADDITIONAL_FEE_LINE}}" in cell.text and not additional_fee_line):
                        rows_to_remove.append((table._tbl, row._tr))
                        break
        
        # Remove the rows
        for tbl, tr in rows_to_remove:
            tbl.remove(tr)
        # Remove paragraphs containing empty fee line placeholders BEFORE replacement
        # This preserves intentional spacing while removing only unused fee lines
        paragraphs_to_remove = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            # Only remove if it contains a fee line placeholder that will be empty
            if ("{{FEE_LINE_2}}" in text and not fee_line_2) or \
               ("{{FEE_LINE_3}}" in text and not fee_line_3) or \
               ("{{ADDITIONAL_FEE_LINE}}" in text and not additional_fee_line):
                paragraphs_to_remove.append(paragraph)
        
        # Remove the specific fee line paragraphs
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

        fill_invoice_template(doc, replacements)
        
        # Add property fees as dynamic rows if they exist
        # This is tricky with python-docx if we don't have a specific placeholder row to clone.
        # For now, we will just ensure the total is correct. 
        # If the user wants property fees listed, we'd need a more complex template logic.
        # Given the constraints, let's assume the "Additional Fee" or similar might be used, 
        # OR we just accept that they are summed in the total but not itemized unless we add more logic.
        # WAIT: The user asked for "ability to add a fee when adding a property".
        # Ideally this should be itemized.
        # Let's try to append them to the table if possible, or just leave as is for now and verify total.
        # Since I can't easily clone rows without a reference, I'll stick to the total for now.

        # Calculate street name (remove number)
        address_parts = customer.property_address.split(' ', 1)
        if len(address_parts) > 1:
            street_name = address_parts[1]
        else:
            street_name = customer.property_address
        
        # Sanitize filename
        safe_period = period_label.replace(' ', '_').replace('/', '-')
        safe_street = street_name.replace(' ', '_').replace('/', '-')
        
        filename = f"Invoice_{safe_period}_{safe_street}.docx"
        
        if return_buffer:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return filename, buffer, total_amount
        else:
            output_path = os.path.join(OUTPUT_DIR, filename)
            doc.save(output_path)
            return filename, output_path, total_amount

    except Exception as e:
        print(f"Error generating invoice: {e}")
        raise e

def generate_invoice_with_template(customer, invoice_date, template_name, **kwargs):
    """Generate invoice and save to database (for manual generation via UI)."""
    session = SessionLocal()
    try:
        period_label = get_period_label(invoice_date, customer.cadence)
        amount = customer.rate 
        start_date, end_date = get_period_dates(invoice_date, customer.cadence)
        period_dates = f"{start_date.strftime('%m/%d/%Y')} - {end_date.strftime('%m/%d/%Y')}"
        
        # Generate invoice in-memory
        filename, buffer, total_amount = _generate_invoice_logic(customer, invoice_date, period_label, period_dates, amount, **kwargs)
        
        # Get sender info from settings
        settings = session.query(Settings).first()
        sender_name = settings.sender_name if settings else "Property Manager"
        
        # Create email content
        fee_type_text = getattr(customer, "fee_type", "Management Fee") or "Management Fee"
        subject = f"Invoice – {period_label} – {customer.property_address}"
        body = (
            f"Hi {customer.name},\n\n"
            f"Attached is your invoice for {period_label} ({fee_type_text}) for the property at {customer.property_address}.\n\n"
            f"Amount due: ${total_amount:,.2f}\n\n"
            f"Thank you,\n{sender_name}"
        )
        
        # Save to database
        invoice_record = Invoice(
            customer_id=customer.id,
            invoice_date=invoice_date,
            period_label=period_label,
            amount=customer.rate, # Store BASE amount (rate) so regeneration works correctly
            # WAIT: If we store total_amount here, then future regenerations might double count fees if we add fees to it again?
            # The Invoice model has 'amount'. If we store total here, we should be careful.
            # But the invoice generation logic takes 'amount' as input.
            # If we pass total_amount back into logic, it will add fees ON TOP of total.
            # So we should probably store the BASE amount in the DB if we want to regenerate with dynamic fees.
            # OR we store the total and don't add fees again?
            # The current logic passes 'customer.rate' (base) to logic.
            # So 'invoice_record.amount' should probably be the TOTAL for display purposes in the list?
            # Let's check how 'amount' is used.
            # In 'generate_invoice_buffer', it uses 'invoice.amount'.
            # If 'invoice.amount' is TOTAL, and we pass it to logic, logic adds fees.
            # So logic would calculate: TOTAL + fees. That's wrong.
            # So 'invoice.amount' MUST be the BASE amount.
            # BUT the user wants the email to say the TOTAL.
            # So we use 'total_amount' for the email body, but store 'amount' (base) in the DB?
            # The 'Invoices' list shows 'inv.amount'. If that shows base, it's confusing.
            # Ideally 'Invoice' model should have 'base_amount' and 'total_amount'.
            # Or we change logic to not add fees if they are already included? No, that's messy.
            # Let's keep 'amount' as BASE amount in DB to preserve regeneration logic.
            # And just use 'total_amount' for the email body.
            # AND maybe we should update the 'Invoices' list to calculate total on the fly?
            # Or add a 'total_amount' column to Invoice?
            # For now, I will just fix the EMAIL as requested.
            # The user complained about the email body.
            file_path=filename,
            email_subject=subject,
            email_body=body,
            # Save extra fees if provided
            fee_2_type=kwargs.get("fee_2_type"),
            fee_2_amount=kwargs.get("fee_2_amount"),
            fee_3_type=kwargs.get("fee_3_type"),
            fee_3_amount=kwargs.get("fee_3_amount"),
            additional_fee_desc=kwargs.get("additional_fee_desc"),
            additional_fee_amount=kwargs.get("additional_fee_amount")
        )
        session.add(invoice_record)
        session.commit()
        
        return invoice_record
    finally:
        session.close()

def generate_invoice_for_customer(customer, invoice_date):
    period_label = get_period_label(invoice_date, customer.cadence)
    start_date, end_date = get_period_dates(invoice_date, customer.cadence)
    period_dates = f"{start_date.strftime('%m/%d/%Y')} - {end_date.strftime('%m/%d/%Y')}"
    amount = customer.rate
    
    # Generate invoice in-memory (don't write to disk - Vercel is read-only)
    filename, buffer, total_amount = _generate_invoice_logic(customer, invoice_date, period_label, period_dates, amount)

    # Get sender info from settings
    session = SessionLocal()
    settings = session.query(Settings).first()
    sender_name = settings.sender_name if settings else "Property Manager"
    session.close()

    fee_type_text = getattr(customer, "fee_type", "Management Fee") or "Management Fee"
    subject = f"Invoice – {period_label} – {customer.property_address}"
    body = (
        f"Hi {customer.name},\n\n"
        f"Attached is your invoice for {period_label} ({fee_type_text}) for the property at {customer.property_address}.\n\n"
        f"Amount due: ${total_amount:,.2f}\n\n"
        f"Thank you,\n{sender_name}"
    )

    invoice = Invoice(
        customer_id=customer.id,
        invoice_date=invoice_date,
        period_label=period_label,
        amount=amount, # Keep as base amount
        file_path=filename, # Store filename only for cloud compatibility
        email_subject=subject,
        email_body=body,
        # Save fee details so they persist for regeneration
        fee_2_type=customer.fee_2_type,
        fee_2_amount=customer.fee_2_rate,
        fee_3_type=customer.fee_3_type,
        fee_3_amount=customer.fee_3_rate,
        additional_fee_desc=customer.additional_fee_desc,
        additional_fee_amount=customer.additional_fee_amount
    )
    
    session = SessionLocal()
    session.add(invoice)
    session.commit()
    session.close()
    
    return invoice

def generate_invoice_buffer(invoice):
    """
    Regenerates the invoice document in-memory for a given Invoice record.
    """
    session = SessionLocal()
    customer = session.query(Customer).get(invoice.customer_id)
    
    # Eagerly load properties to avoid lazy loading issues after session close
    if customer and customer.properties:
        _ = list(customer.properties)  # Force load
    
    session.close()
    
    if not customer:
        raise ValueError("Customer not found")
        
    # Reconstruct parameters
    # Note: In a real app, we might want to store period_dates in the Invoice model too.
    # For now, we recalculate them based on the invoice date and customer cadence.
    # This assumes the cadence hasn't changed in a way that affects the past invoice period logic.
    
    start_date, end_date = get_period_dates(invoice.invoice_date, customer.cadence)
    period_dates = f"{start_date.strftime('%m/%d/%Y')} - {end_date.strftime('%m/%d/%Y')}"
    
    filename, buffer, _ = _generate_invoice_logic(
        customer, 
        invoice.invoice_date, 
        invoice.period_label, 
        period_dates, 
        invoice.amount, 
        return_buffer=True,
        fee_2_type=invoice.fee_2_type,
        fee_2_amount=invoice.fee_2_amount,
        fee_3_type=invoice.fee_3_type,
        fee_3_amount=invoice.fee_3_amount,
        additional_fee_desc=invoice.additional_fee_desc,
        additional_fee_amount=invoice.additional_fee_amount
    )
    return filename, buffer
