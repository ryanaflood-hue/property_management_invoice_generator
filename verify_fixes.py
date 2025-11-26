import requests
import os
from docx import Document
from datetime import date

BASE_URL = "http://127.0.0.1:5000"

def verify_fixes():
    print("--- Starting Verification ---")
    
    # 0. Health Check
    print(f"Checking {BASE_URL}/...")
    try:
        resp = requests.get(f"{BASE_URL}/")
        print(f"Health Check: {resp.status_code}")
    except Exception as e:
        print(f"Health Check Failed: {e}")
        return

    # 1. Add Fee Types
    print("1. Adding Fee Types...")
    requests.post(f"{BASE_URL}/settings/fee-types", data={"name": "Management Fee"})
    requests.post(f"{BASE_URL}/settings/fee-types", data={"name": "Late Fee"})
    
    # 2. Create Customer
    print("2. Creating Customer...")
    customer_data = {
        "name": "Verification Customer",
        "email": "verify@example.com",
        "property_address": "123 Verify St",
        "property_city": "Verify City",
        "property_state": "VS",
        "property_zip": "99999",
        "rate": "100.00",
        "cadence": "monthly",
        "fee_type": "Management Fee",
        "next_bill_date": "2025-10-01",
        "fee_2_type": "Late Fee",
        "fee_2_rate": "50.00", # Default is 50
        # Fee 3 is empty by default
    }
    url = f"{BASE_URL}/customers/new"
    print(f"POSTing to {url}")
    resp = requests.post(url, data=customer_data)
    if resp.status_code != 200 and resp.status_code != 302:
        print(f"[FAIL] Failed to create customer: {resp.status_code}")
        print(resp.text)
        return

    # 3. Generate Invoice with Overrides
    print("3. Generating Invoice with Overrides...")
    # We need the customer ID. Assuming it's the last one created or we can query.
    # Since we just created it, and DB was fresh, it should be ID 1.
    # But let's check the list to be sure? No, just assume 1 for now.
    
    invoice_data = {
        "customer_id": "1",
        "invoice_date": "2025-11-24",
        "template_name": "base_invoice_template.docx",
        "fee_2_type": "Late Fee",
        "fee_2_amount": "99.00", # OVERRIDE: Should be 99, not 50
        "fee_3_type": "Late Fee",
        "fee_3_amount": "75.00", # NEW FEE: Should appear
        "additional_fee_desc": "Air Purifier",
        "additional_fee_amount": "300.00"
    }
    
    resp = requests.post(f"{BASE_URL}/generate-invoice", data=invoice_data)
    if resp.status_code != 200:
        print(f"[FAIL] Failed to generate invoice: {resp.status_code}")
        print(resp.text)
        return
        
    print("[OK] Invoice generated successfully.")
    
    # 4. Find the generated file
    # It should be in generated_invoices
    output_dir = "generated_invoices"
    files = sorted([os.path.join(output_dir, f) for f in os.listdir(output_dir)], key=os.path.getmtime)
    if not files:
        print("[FAIL] No invoice found in generated_invoices.")
        return
        
    latest_file = files[-1]
    print(f"Inspecting: {latest_file}")
    
    # 5. Inspect Content
    doc = Document(latest_file)
    text = "\n".join([p.text for p in doc.paragraphs])
    
    print("\n[Content Check]")
    
    # Check Fee 2 Override
    if "$99.00" in text:
        print("[OK] Fee 2 Override ($99.00) FOUND.")
    elif "$50.00" in text:
        print("[FAIL] Fee 2 Override FAILED. Found default $50.00.")
    else:
        print("[FAIL] Fee 2 Value NOT FOUND.")
        
    # Check Fee 3
    if "$75.00" in text:
        print("[OK] Fee 3 ($75.00) FOUND.")
    else:
        print("[FAIL] Fee 3 ($75.00) MISSING.")
        
    # Check Total
    # 100 + 99 + 75 + 300 = 574
    if "$574.00" in text:
        print("[OK] Total ($574.00) MATCHES expected.")
    else:
        print("[FAIL] Total does NOT match expected $574.00.")
        # Find what total is
        for p in doc.paragraphs:
            if "Total due:" in p.text:
                print(f"   Actual Total: {p.text}")

if __name__ == "__main__":
    verify_fixes()
