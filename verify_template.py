from docx import Document
import os

def verify_template():
    path = "invoice_templates/base_invoice_template.docx"
    if not os.path.exists(path):
        print(f"Error: {path} not found.")
        return

    doc = Document(path)
    
    print(f"--- Inspecting {path} ---")
    
    placeholders = [
        "{{FEE_LINE_2}}",
        "{{FEE_LINE_3}}",
        "{{ADDITIONAL_FEE_LINE}}",
        "{{TOTAL_AMOUNT}}"
    ]
    
    found = {p: False for p in placeholders}
    
    print("\n[Paragraphs]")
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        for ph in placeholders:
            if ph in text:
                found[ph] = True
                print(f"Found {ph} in paragraph {i}: '{text}'")
            # Check for split runs
            elif ph.replace("{", "").replace("}", "") in text and "{" in text:
                 print(f"WARNING: Possible split placeholder in paragraph {i}: '{text}'")

    print("\n[Tables]")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    text = p.text
                    for ph in placeholders:
                        if ph in text:
                            found[ph] = True
                            print(f"Found {ph} in Table {t_idx}, Row {r_idx}, Cell {c_idx}: '{text}'")

    print("\n--- Summary ---")
    for ph, is_found in found.items():
        status = "✅ FOUND" if is_found else "❌ MISSING"
        print(f"{ph}: {status}")

if __name__ == "__main__":
    verify_template()
