import os
from docx import Document
import re

GENERATED_DIR = r"C:\Development\invoice_automation\generated_invoices"
PLACEHOLDER_PATTERN = re.compile(r"{{(.*?)}}")

def check_invoices():
    files = [f for f in os.listdir(GENERATED_DIR) if f.endswith('.docx')]
    files.sort(key=lambda x: os.path.getmtime(os.path.join(GENERATED_DIR, x)), reverse=True)
    
    # Check the last 2 files
    for f in files[:2]:
        path = os.path.join(GENERATED_DIR, f)
        print(f"Checking {f}...")
        doc = Document(path)
        found = set()
        for para in doc.paragraphs:
            found.update(PLACEHOLDER_PATTERN.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found.update(PLACEHOLDER_PATTERN.findall(cell.text))
        
        if found:
            print(f"  Missing values: {found}")
        else:
            print("  No missing values found.")

if __name__ == "__main__":
    check_invoices()
