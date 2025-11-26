from docx import Document
import re

TEMPLATE_PATH = r"C:\Development\invoice_automation\invoice_templates\base_invoice_template.docx"

PLACEHOLDER_PATTERN = re.compile(r"{{(.*?)}}")

def extract_placeholders(docx_path):
    doc = Document(docx_path)
    found = set()
    for para in doc.paragraphs:
        found.update(PLACEHOLDER_PATTERN.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found.update(PLACEHOLDER_PATTERN.findall(cell.text))
    return sorted(found)

if __name__ == "__main__":
    placeholders = extract_placeholders(TEMPLATE_PATH)
    print("Found placeholders:")
    for ph in placeholders:
        print(f"{{{{{ph}}}}}")
