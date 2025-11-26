from invoice_generator import _generate_invoice_logic
from models import Customer, Property
from datetime import date
from docx import Document
import os

def reproduce():
    # Setup Customer
    c = Customer(
        id=1,
        name="Test Customer",
        email="test@example.com",
        property_address="123 Test St",
        property_city="Test City",
        property_state="TS",
        property_zip="12345",
        rate=100.0,
        cadence="monthly",
        fee_type="Management Fee",
        # Default fees (to test override)
        fee_2_type="Default Fee 2",
        fee_2_rate=50.0,
        fee_3_type="Default Fee 3",
        fee_3_rate=25.0,
        additional_fee_desc="Default Add Fee",
        additional_fee_amount=10.0,
        next_bill_date=date(2025, 10, 1)
    )
    # Add Property with fee
    p = Property(address="123 Test St", fee_amount=125.0)
    c.properties = [p]

    # Manual Overrides (matching user scenario)
    kwargs = {
        "fee_2_type": "Late Fee",
        "fee_2_amount": 50.0,
        "fee_3_type": "Release Fee",
        "fee_3_amount": 30.0,
        "additional_fee_desc": "Air Purifier Fee",
        "additional_fee_amount": 300.0
    }

    print("Generating invoice...")
    # Pass return_buffer=False to save file
    filename, path, total_amount = _generate_invoice_logic(
        c,
        date(2025, 11, 24),
        "4th quarter 2025",
        "10/01/2025 - 12/31/2025",
        100.0,
        return_buffer=False,
        **kwargs
    )
    
    print(f"Generated: {path}")
    print(f"Total Amount Returned: ${total_amount:,.2f}")
    
    if total_amount == 605.00:
        print("[OK] Total Amount matches expected $605.00")
    else:
        print(f"[FAIL] Total Amount {total_amount} does not match expected $605.00")
    
    # Inspect the generated file
    doc = Document(path)
    print("\n[Generated Content]")
    found_fee_3 = False
    found_prop_fee = False
    
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"P: '{p.text}'")
        if "30.00" in p.text:
            found_fee_3 = True
        if "125.00" in p.text:
            found_prop_fee = True
            
    if found_fee_3:
        print("\n[OK] Fee 3 ($30.00) FOUND in output.")
    else:
        print("\n[FAIL] Fee 3 ($30.00) MISSING from output.")

    if found_prop_fee:
        print("\n[OK] Property Fee ($125.00) FOUND in output.")
    else:
        print("\n[FAIL] Property Fee ($125.00) MISSING from output.")

    # Check formatting for Additional Fee Line
    print("\n[Formatting Check]")
    for p in doc.paragraphs:
        if "Air Purifier Fee" in p.text:
            space_after = p.paragraph_format.space_after
            line_spacing = p.paragraph_format.line_spacing
            print(f"Additional Fee Line - Space After: {space_after}, Line Spacing: {line_spacing}")
            if space_after == 0 or (hasattr(space_after, 'pt') and space_after.pt == 0.0):
                print("[OK] Space After is 0 (Tight Spacing)")
            else:
                print(f"[FAIL] Space After is {space_after} (Not Tight)")

if __name__ == "__main__":
    reproduce()
