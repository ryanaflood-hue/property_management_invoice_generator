from docx import Document
import os

TEMPLATE_PATH = os.path.join("invoice_templates", "base_invoice_template.docx")

def inspect_template():
    if not os.path.exists(TEMPLATE_PATH):
        print(f"Template not found at {TEMPLATE_PATH}")
        return

    doc = Document(TEMPLATE_PATH)
    
    placeholders = ["{{FEE_LINE_2}}", "{{FEE_LINE_3}}", "{{ADDITIONAL_FEE_LINE}}"]
    found = {p: False for p in placeholders}
    
    print("Searching paragraphs...")
    for p in doc.paragraphs:
        for ph in placeholders:
            if ph in p.text:
                found[ph] = True
                print(f"Found {ph} in paragraph: {p.text.strip()}")

    print("\nSearching tables...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    for ph in placeholders:
                        if ph in p.text:
                            found[ph] = True
                            print(f"Found {ph} in Table {table_idx}, Row {row_idx}: {p.text.strip()}")

    print("\nResults:")
    for ph, was_found in found.items():
        print(f"{ph}: {'FOUND' if was_found else 'MISSING'}")

if __name__ == "__main__":
    inspect_template()
