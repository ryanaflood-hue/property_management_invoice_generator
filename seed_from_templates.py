import os
import re
from docx import Document
from models import SessionLocal, Customer, FeeType, init_db
from datetime import date

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "invoice_templates")

def extract_money(text):
    # Find all money patterns like $150 or $1,200.50
    matches = re.findall(r'\$\s?([0-9,]+(?:\.[0-9]{2})?)', text)
    if matches:
        # Return the last one found in the line, assuming it's the total for that line
        return float(matches[-1].replace(',', ''))
    return 0.0

def parse_address(full_address):
    """Parse full address into street, city, state, zip components."""
    # Common patterns:
    # "1214 S 115th ST, West Allis, WI  53214"
    # "N62W12921 River Heights Dr"
    # "2085 Le Jardin Ct."
    
    street = ""
    city = ""
    state = ""
    zip_code = ""
    
    # Try to find zip code (5 digits)
    zip_match = re.search(r'\b(\d{5})\b', full_address)
    if zip_match:
        zip_code = zip_match.group(1)
        # Remove zip from address
        full_address = full_address.replace(zip_code, '').strip()
    
    # Try to find state (2 uppercase letters)
    state_match = re.search(r'\b([A-Z]{2})\b', full_address)
    if state_match:
        state = state_match.group(1)
        # Remove state from address
        full_address = full_address.replace(state, '').strip()
    
    # Split by comma to separate street from city
    parts = [p.strip() for p in full_address.split(',')]
    
    if len(parts) >= 2:
        street = parts[0]
        city = parts[1]
    elif len(parts) == 1:
        street = parts[0]
    
    # Clean up trailing commas
    street = street.rstrip(',').strip()
    city = city.rstrip(',').strip()
    
    return street, city, state, zip_code

def seed_customers():
    print("Initializing DB...")
    init_db()
    session = SessionLocal()
    
    # Ensure Management Fee exists
    if not session.query(FeeType).filter_by(name="Management Fee").first():
        session.add(FeeType(name="Management Fee"))
        session.commit()

    files = [f for f in os.listdir(TEMPLATE_DIR) if f.lower().endswith('.docx')]
    
    count = 0
    for f in files:
        if f == "base_invoice_template.docx" or f.startswith("~"):
            continue
            
        path = os.path.join(TEMPLATE_DIR, f)
        print(f"Processing {f}...")
        
        try:
            doc = Document(path)
            
            name = ""
            address = ""
            email = ""
            rate = 0.0
            cadence = "monthly" # Default
            fee_type = "Management Fee"
            
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                
                # Extract Name
                if text.upper().startswith("TO:"):
                    # Try to get name from same line
                    parts = text.split(":", 1)
                    if len(parts) > 1 and parts[1].strip():
                        name = parts[1].strip()
                    else:
                        # Name might be on next line, but simple parsing is hard here.
                        # Let's assume it's on the same line for now based on analysis.
                        pass
                
                # Extract Address
                if text.upper().startswith("FOR:"):
                    parts = text.split(":", 1)
                    if len(parts) > 1 and parts[1].strip():
                        address = parts[1].strip()
                
                # Extract Email
                email_matches = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
                for email_match in email_matches:
                    # Filter out sender emails
                    if "linda" not in email_match.lower() and "stonegate" not in email_match.lower():
                         # Assume the first non-sender email is the customer's
                         if not email or email == "change@me.com":
                             email = email_match

                # Extract Rate & Cadence from management line
                if "management" in text.lower() or "quarter" in text.lower():
                    line_rate = extract_money(text)
                    if line_rate > 0:
                        rate = line_rate
                    
                    if "quarter" in text.lower():
                        cadence = "quarterly"
                    elif "year" in text.lower() or "annual" in text.lower():
                        cadence = "yearly"
            
            # Fallback for rate if not found in management line
            if rate == 0:
                for p in doc.paragraphs:
                    if "Total due" in p.text:
                        rate = extract_money(p.text)
                        break

            if name and address:
                # Parse address into components
                street, city, state, zip_code = parse_address(address)
                
               # Check if exists
                existing = session.query(Customer).filter(Customer.name == name).first()
                if not existing:
                    print(f"  -> Adding {name} ({street}, {city}, {state} {zip_code}) - ${rate} {cadence}")
                    c = Customer(
                        name=name,
                        email=email if email else "change@me.com",
                        property_address=street,
                        property_city=city,
                        property_state=state,
                        property_zip=zip_code,
                        rate=rate,
                        cadence=cadence,
                        fee_type=fee_type,
                        next_bill_date=date(2025, 10, 1)
                    )
                    session.add(c)
                    count += 1
                else:
                    print(f"  -> Updating existing customer: {name}")
                    # Force update all fields to match template
                    existing.property_address = street
                    existing.property_city = city
                    existing.property_state = state
                    existing.property_zip = zip_code
                    existing.rate = rate
                    existing.cadence = cadence
                    existing.fee_type = fee_type
                    # Reset next_bill_date to today so "Run Daily Batch" picks them up
                    existing.next_bill_date = date(2025, 10, 1)
                    
                    if email:
                        existing.email = email
                    
                    session.add(existing)
                    count += 1
            else:
                print(f"  -> Could not extract Name or Address from {f}")

        except Exception as e:
            print(f"  -> Error processing {f}: {e}")

    session.commit()
    session.close()
    print(f"Done! Added {count} new customers.")

if __name__ == "__main__":
    from datetime import date
    seed_customers()
