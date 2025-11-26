import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "invoice_templates")
SAMPLE_FILE = "Comm. inv. 2025 2nd quarter 115th ST.docx"

def analyze_docx():
    path = os.path.join(TEMPLATE_DIR, SAMPLE_FILE)
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return

    doc = Document(path)
    
    print(f"--- Analyzing {SAMPLE_FILE} ---")
    print("PARAGRAPHS:")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"{i}: {p.text}")
            
    print("\nTABLES:")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"  {row_text}")

if __name__ == "__main__":
    analyze_docx()
